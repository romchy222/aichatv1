"""
File processing utilities for handling various file types:
- Images (JPG, PNG, etc.)
- Documents (DOCX, PDF, TXT)
- Spreadsheets (XLSX, CSV)
"""

import os
import json
import logging
from io import BytesIO
from PIL import Image
from docx import Document
import pandas as pd
import PyPDF2

# Optional imports
try:
    import cv2
    import numpy as np
    CV2_AVAILABLE = True
except ImportError:
    CV2_AVAILABLE = False

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False

from django.conf import settings

logger = logging.getLogger('agent')


class FileProcessor:
    """Base class for file processing"""
    
    def __init__(self):
        self.supported_extensions = []
    
    def can_process(self, file_extension):
        """Check if this processor can handle the file"""
        return file_extension.lower() in self.supported_extensions
    
    def process(self, file_path):
        """Process the file and return extracted information"""
        raise NotImplementedError


class ImageProcessor(FileProcessor):
    """Processor for image files"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp', 'tiff']
    
    def process(self, file_path):
        """Process image file and extract information"""
        try:
            # Load image with PIL
            with Image.open(file_path) as img:
                # Basic image info
                info = {
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'mode': img.mode,
                    'size_bytes': os.path.getsize(file_path)
                }
                
                # Try to extract text using OCR (simplified approach)
                extracted_text = self.extract_text_from_image(file_path)
                
                # Analyze image content
                analysis = self.analyze_image_content(file_path)
                
                return {
                    'success': True,
                    'extracted_text': extracted_text,
                    'analysis': {
                        **info,
                        **analysis
                    },
                    'summary': self.generate_image_summary(info, extracted_text, analysis)
                }
                
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'extracted_text': '',
                'analysis': {},
                'summary': f'Ошибка обработки изображения: {str(e)}'
            }
    
    def extract_text_from_image(self, file_path):
        """Extract text from image using OpenCV (basic approach)"""
        try:
            if not CV2_AVAILABLE:
                return "Изображение обнаружено (OCR недоступен)"
            
            # This is a placeholder for OCR functionality
            # In production, you would use pytesseract or similar
            img = cv2.imread(file_path)
            if img is not None:
                # Basic image preprocessing
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                # For now, return basic description
                return "Обнаружено изображение (требуется OCR для извлечения текста)"
            return ""
        except Exception as e:
            logger.warning(f"Could not extract text from image: {e}")
            return ""
    
    def analyze_image_content(self, file_path):
        """Analyze image content and detect features"""
        try:
            if not CV2_AVAILABLE:
                return {'note': 'Расширенный анализ изображений недоступен'}
            
            img = cv2.imread(file_path)
            if img is None:
                return {}
            
            # Basic image analysis
            height, width, channels = img.shape
            
            # Color analysis
            mean_color = np.mean(img, axis=(0, 1))
            
            # Brightness analysis
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            brightness = np.mean(gray)
            
            return {
                'dimensions': f"{width}x{height}",
                'channels': channels,
                'brightness': float(brightness),
                'dominant_colors': [float(c) for c in mean_color],
                'is_dark': brightness < 100,
                'is_bright': brightness > 200
            }
        except Exception as e:
            logger.warning(f"Could not analyze image content: {e}")
            return {}
    
    def generate_image_summary(self, info, text, analysis):
        """Generate human-readable summary of image"""
        summary = f"Изображение {info['format']} размером {info['width']}x{info['height']} пикселей"
        
        if analysis.get('is_dark'):
            summary += " (темное изображение)"
        elif analysis.get('is_bright'):
            summary += " (яркое изображение)"
        
        if text and len(text) > 10:
            summary += f". Содержит текст: {text[:100]}..."
        
        return summary


class DocumentProcessor(FileProcessor):
    """Processor for document files"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['docx', 'doc', 'txt', 'rtf']
    
    def process(self, file_path):
        """Process document file and extract text"""
        try:
            file_extension = file_path.split('.')[-1].lower()
            
            if file_extension == 'docx':
                return self.process_docx(file_path)
            elif file_extension == 'txt':
                return self.process_txt(file_path)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported document format: {file_extension}',
                    'extracted_text': '',
                    'analysis': {},
                    'summary': f'Неподдерживаемый формат документа: {file_extension}'
                }
                
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'extracted_text': '',
                'analysis': {},
                'summary': f'Ошибка обработки документа: {str(e)}'
            }
    
    def process_docx(self, file_path):
        """Process DOCX file"""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            full_text = '\n'.join(text_content)
            
            # Extract text from tables
            table_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_content.append(' | '.join(row_text))
            
            if table_content:
                full_text += '\n\nТаблицы:\n' + '\n'.join(table_content)
            
            # Basic analysis
            analysis = {
                'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
                'table_count': len(doc.tables),
                'word_count': len(full_text.split()),
                'character_count': len(full_text),
                'file_size': os.path.getsize(file_path)
            }
            
            summary = f"Документ Word с {analysis['paragraph_count']} абзацами"
            if analysis['table_count'] > 0:
                summary += f" и {analysis['table_count']} таблицами"
            summary += f". Всего слов: {analysis['word_count']}"
            
            return {
                'success': True,
                'extracted_text': full_text,
                'analysis': analysis,
                'summary': summary
            }
            
        except Exception as e:
            raise Exception(f"Error processing DOCX: {e}")
    
    def process_txt(self, file_path):
        """Process TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            lines = content.split('\n')
            words = content.split()
            
            analysis = {
                'line_count': len(lines),
                'word_count': len(words),
                'character_count': len(content),
                'file_size': os.path.getsize(file_path)
            }
            
            summary = f"Текстовый файл с {analysis['line_count']} строками и {analysis['word_count']} словами"
            
            return {
                'success': True,
                'extracted_text': content,
                'analysis': analysis,
                'summary': summary
            }
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['cp1251', 'iso-8859-1', 'cp866']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    return self.process_txt_with_content(content, file_path)
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with any common encoding")
    
    def process_txt_with_content(self, content, file_path):
        """Process text content"""
        lines = content.split('\n')
        words = content.split()
        
        analysis = {
            'line_count': len(lines),
            'word_count': len(words),
            'character_count': len(content),
            'file_size': os.path.getsize(file_path)
        }
        
        summary = f"Текстовый файл с {analysis['line_count']} строками и {analysis['word_count']} словами"
        
        return {
            'success': True,
            'extracted_text': content,
            'analysis': analysis,
            'summary': summary
        }


class SpreadsheetProcessor(FileProcessor):
    """Processor for spreadsheet files"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['xlsx', 'xls', 'csv', 'ods']
    
    def process(self, file_path):
        """Process spreadsheet file"""
        try:
            file_extension = file_path.split('.')[-1].lower()
            
            if file_extension in ['xlsx', 'xls']:
                return self.process_excel(file_path)
            elif file_extension == 'csv':
                return self.process_csv(file_path)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported spreadsheet format: {file_extension}',
                    'extracted_text': '',
                    'analysis': {},
                    'summary': f'Неподдерживаемый формат таблицы: {file_extension}'
                }
                
        except Exception as e:
            logger.error(f"Error processing spreadsheet {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'extracted_text': '',
                'analysis': {},
                'summary': f'Ошибка обработки таблицы: {str(e)}'
            }
    
    def process_excel(self, file_path):
        """Process Excel file"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            all_text = []
            total_rows = 0
            total_cols = 0
            
            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert DataFrame to text
                sheet_text = f"\n=== Лист: {sheet_name} ===\n"
                
                # Add column headers
                headers = list(df.columns)
                sheet_text += " | ".join(str(h) for h in headers) + "\n"
                sheet_text += "-" * (len(headers) * 15) + "\n"
                
                # Add data rows (limit to first 100 rows for performance)
                for idx, row in df.head(100).iterrows():
                    row_text = " | ".join(str(cell) if pd.notna(cell) else "" for cell in row)
                    sheet_text += row_text + "\n"
                
                if len(df) > 100:
                    sheet_text += f"\n... (показано 100 из {len(df)} строк)\n"
                
                all_text.append(sheet_text)
                total_rows += len(df)
                total_cols = max(total_cols, len(df.columns))
            
            full_text = '\n'.join(all_text)
            
            analysis = {
                'sheet_count': len(sheet_names),
                'total_rows': total_rows,
                'max_columns': total_cols,
                'sheet_names': sheet_names,
                'file_size': os.path.getsize(file_path)
            }
            
            summary = f"Excel файл с {analysis['sheet_count']} листами, всего {analysis['total_rows']} строк"
            
            return {
                'success': True,
                'extracted_text': full_text,
                'analysis': analysis,
                'summary': summary
            }
            
        except Exception as e:
            raise Exception(f"Error processing Excel: {e}")
    
    def process_csv(self, file_path):
        """Process CSV file"""
        try:
            # Try different encodings and separators
            for encoding in ['utf-8', 'cp1251', 'iso-8859-1']:
                for sep in [',', ';', '\t']:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding, sep=sep)
                        if len(df.columns) > 1:  # Valid CSV found
                            break
                    except:
                        continue
                else:
                    continue
                break
            else:
                raise Exception("Could not parse CSV file")
            
            # Convert to text
            text_content = []
            
            # Add headers
            headers = list(df.columns)
            text_content.append(" | ".join(str(h) for h in headers))
            text_content.append("-" * (len(headers) * 15))
            
            # Add data (limit to first 100 rows)
            for idx, row in df.head(100).iterrows():
                row_text = " | ".join(str(cell) if pd.notna(cell) else "" for cell in row)
                text_content.append(row_text)
            
            if len(df) > 100:
                text_content.append(f"\n... (показано 100 из {len(df)} строк)")
            
            full_text = '\n'.join(text_content)
            
            analysis = {
                'row_count': len(df),
                'column_count': len(df.columns),
                'column_names': headers,
                'file_size': os.path.getsize(file_path)
            }
            
            summary = f"CSV файл с {analysis['row_count']} строками и {analysis['column_count']} столбцами"
            
            return {
                'success': True,
                'extracted_text': full_text,
                'analysis': analysis,
                'summary': summary
            }
            
        except Exception as e:
            raise Exception(f"Error processing CSV: {e}")


class PDFProcessor(FileProcessor):
    """Processor for PDF files"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['pdf']
    
    def process(self, file_path):
        """Process PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from all pages
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"=== Страница {page_num + 1} ===\n{page_text}")
                    except Exception as e:
                        text_content.append(f"=== Страница {page_num + 1} ===\n[Ошибка извлечения текста: {e}]")
                
                full_text = '\n\n'.join(text_content)
                
                analysis = {
                    'page_count': len(pdf_reader.pages),
                    'word_count': len(full_text.split()),
                    'character_count': len(full_text),
                    'file_size': os.path.getsize(file_path),
                    'has_metadata': bool(pdf_reader.metadata)
                }
                
                # Add metadata if available
                if pdf_reader.metadata:
                    metadata = {}
                    for key, value in pdf_reader.metadata.items():
                        if value:
                            metadata[key] = str(value)
                    analysis['metadata'] = metadata
                
                summary = f"PDF документ с {analysis['page_count']} страницами, {analysis['word_count']} слов"
                
                return {
                    'success': True,
                    'extracted_text': full_text,
                    'analysis': analysis,
                    'summary': summary
                }
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'extracted_text': '',
                'analysis': {},
                'summary': f'Ошибка обработки PDF: {str(e)}'
            }


class FileProcessorManager:
    """Manager for all file processors"""
    
    def __init__(self):
        self.processors = {
            'image': ImageProcessor(),
            'document': DocumentProcessor(),
            'spreadsheet': SpreadsheetProcessor(),
            'pdf': PDFProcessor()
        }
    
    def get_file_type(self, filename):
        """Determine file type based on extension"""
        extension = filename.split('.')[-1].lower() if '.' in filename else ''
        
        for file_type, processor in self.processors.items():
            if processor.can_process(extension):
                return file_type
        
        return 'other'
    
    def process_file(self, file_path, file_type=None):
        """Process file using appropriate processor"""
        if file_type is None:
            filename = os.path.basename(file_path)
            file_type = self.get_file_type(filename)
        
        if file_type == 'other':
            return {
                'success': False,
                'error': 'Unsupported file type',
                'extracted_text': '',
                'analysis': {},
                'summary': 'Неподдерживаемый тип файла'
            }
        
        processor = self.processors.get(file_type)
        if not processor:
            return {
                'success': False,
                'error': f'No processor for file type: {file_type}',
                'extracted_text': '',
                'analysis': {},
                'summary': f'Нет обработчика для типа файла: {file_type}'
            }
        
        return processor.process(file_path)
    
    def get_mime_type(self, file_path):
        """Get MIME type of file"""
        try:
            if MAGIC_AVAILABLE:
                mime_type = magic.from_file(file_path, mime=True)
                return mime_type
            else:
                # Fallback to basic extension mapping
                extension = file_path.split('.')[-1].lower()
                mime_map = {
                    'jpg': 'image/jpeg', 'jpeg': 'image/jpeg', 'png': 'image/png',
                    'pdf': 'application/pdf', 'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    'csv': 'text/csv', 'txt': 'text/plain'
                }
                return mime_map.get(extension, 'application/octet-stream')
        except Exception as e:
            logger.warning(f"Could not determine MIME type for {file_path}: {e}")
            return 'application/octet-stream'